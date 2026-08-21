#!/usr/bin/env python3
"""
resume_template.py — Shared PDF/DOCX formatting for all resume generator scripts.

All agents (Wibey, OpenAI, Claude, OpenRouter) produce scripts that import this
module, ensuring consistent formatting regardless of which AI generated the content.

Usage in generated scripts:
    from resume_template import generate_docx, generate_pdf, sanitize
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


# ── Colors ───────────────────────────────────────────────────────────────────
DARK_BLUE_RGB = RGBColor(0x1B, 0x3A, 0x5C)
MED_BLUE_RGB = RGBColor(0x2C, 0x5F, 0x8A)
GRAY_RGB = RGBColor(0x55, 0x55, 0x55)

DARK_BLUE = (27, 58, 92)
MED_BLUE = (44, 95, 138)
GRAY = (100, 100, 100)
BLACK = (0, 0, 0)


# ── Sanitize ─────────────────────────────────────────────────────────────────
def sanitize(text):
    """Replace Unicode characters with ASCII equivalents for PDF core fonts."""
    return (
        text.replace("–", "-")   # en dash
        .replace("—", "-")       # em dash
        .replace("‘", "'")       # left single quote
        .replace("’", "'")       # right single quote
        .replace("“", '"')       # left double quote
        .replace("”", '"')       # right double quote
        .replace("…", "...")     # ellipsis
        .replace("→", "->")     # arrow
        .replace("•", "-")      # bullet
        .replace(" ", " ")      # non-breaking space
    )


# ══════════════════════════════════════════════════════════════════════════════
# DOCX generation
# ══════════════════════════════════════════════════════════════════════════════

def _add_heading_styled(doc, text, level, color=DARK_BLUE_RGB):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    colon_idx = text.find(":")
    if 0 < colon_idx < 80:
        run_bold = p.add_run(text[: colon_idx + 1])
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_rest = p.add_run(text[colon_idx + 1:])
        run_rest.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def _generate_experience_docx(doc, exp):
    if exp.get("company"):
        p_company = doc.add_paragraph()
        run = p_company.add_run(exp["company"])
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_BLUE_RGB
        if exp.get("location"):
            run2 = p_company.add_run(f"  {exp['location']}")
            run2.font.size = Pt(10)
            run2.font.color.rgb = GRAY_RGB

    p_title = doc.add_paragraph()
    if exp.get("title"):
        run = p_title.add_run(exp["title"])
        run.italic = True
        run.font.size = Pt(10)
    run2 = p_title.add_run(f"  {exp['dates']}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY_RGB

    if exp.get("intro"):
        p = doc.add_paragraph()
        run = p.add_run(exp["intro"])
        run.font.size = Pt(10)

    for sub_title, bullets in exp.get("subsections", []):
        if sub_title:
            _add_heading_styled(doc, sub_title, level=3, color=MED_BLUE_RGB)
        for bullet in bullets:
            _add_bullet(doc, bullet)


def generate_docx(data, output_path):
    """Generate a DOCX resume from structured data.

    Args:
        data: dict with keys: name, phone, email, linkedin, location,
              summary, skills (dict), experiences (list of dicts),
              education (list), certificates (list), awards (list), patents (list)
        output_path: path for the output .docx file

    Returns:
        output_path
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Name
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(data["name"])
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = DARK_BLUE_RGB

    # Contact
    contact_parts = [data.get("phone", ""), data.get("email", ""),
                     data.get("linkedin", ""), data.get("location", "")]
    contact_line = "  ".join(p for p in contact_parts if p)
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_p.add_run(contact_line)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_RGB

    # Summary
    _add_heading_styled(doc, "Summary", level=1)
    p = doc.add_paragraph()
    run = p.add_run(data["summary"])
    run.font.size = Pt(10)

    # Skills
    _add_heading_styled(doc, "Skills", level=1)
    for category, skills in data["skills"].items():
        p = doc.add_paragraph(style="List Bullet")
        run_cat = p.add_run(f"{category}: ")
        run_cat.bold = True
        run_cat.font.size = Pt(10)
        run_skills = p.add_run(skills)
        run_skills.font.size = Pt(10)

    # Experience
    _add_heading_styled(doc, "Experience", level=1)
    for exp in data["experiences"]:
        _generate_experience_docx(doc, exp)

    # Education
    _add_heading_styled(doc, "Education", level=1)
    for edu in data.get("education", []):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(edu)
        run.font.size = Pt(10)

    # Certificates
    if data.get("certificates"):
        _add_heading_styled(doc, "Certificates", level=1)
        for cert in data["certificates"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(cert)
            run.font.size = Pt(10)

    # Awards
    if data.get("awards"):
        _add_heading_styled(doc, "Awards", level=1)
        for award in data["awards"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(award)
            run.font.size = Pt(10)

    # Patents
    if data.get("patents"):
        _add_heading_styled(doc, "Patents", level=1)
        for patent in data["patents"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(patent)
            run.font.size = Pt(10)

    doc.save(output_path)
    print(f"Word resume saved to: {output_path}")
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# PDF generation
# ══════════════════════════════════════════════════════════════════════════════

class ResumePDF(FPDF):
    def __init__(self, candidate_name=""):
        super().__init__()
        self._candidate_name = candidate_name

    def header(self):
        pass

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(*GRAY)
        self.cell(0, 10, f"{self._candidate_name} - Page {self.page_no()}/{{nb}}", align="C")

    def section_heading(self, text):
        self.set_font("Helvetica", "B", 13)
        self.set_text_color(*DARK_BLUE)
        self.cell(0, 9, sanitize(text), new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(*MED_BLUE)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def sub_heading(self, text):
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(*MED_BLUE)
        self.multi_cell(0, 5, sanitize(text), new_x="LMARGIN", new_y="NEXT")
        self.ln(1)

    def body_text(self, text):
        self.set_font("Helvetica", "", 9)
        self.set_text_color(*BLACK)
        self.multi_cell(0, 4.5, sanitize(text), new_x="LMARGIN", new_y="NEXT")
        self.ln(1)

    def bullet_point(self, text):
        text = sanitize(text)
        self.set_font("Helvetica", "", 9)
        self.set_text_color(*BLACK)
        indent = self.l_margin + 5
        self.set_x(self.l_margin)
        self.cell(5, 4.5, "-")
        colon_idx = text.find(":")
        if 0 < colon_idx < 80:
            self.set_font("Helvetica", "B", 9)
            bold_part = text[: colon_idx + 1]
            rest_part = text[colon_idx + 1:]
            self.write(4.5, bold_part)
            self.set_font("Helvetica", "", 9)
            self.write(4.5, rest_part)
            self.ln(5)
        else:
            self.multi_cell(self.w - self.r_margin - indent, 4.5, text,
                            new_x="LMARGIN", new_y="NEXT")
        self.ln(0.5)

    def company_line(self, company, location):
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(*DARK_BLUE)
        self.cell(self.w - self.l_margin - self.r_margin - 50, 6, sanitize(company))
        self.set_font("Helvetica", "", 9)
        self.set_text_color(*GRAY)
        self.cell(50, 6, sanitize(location), align="R", new_x="LMARGIN", new_y="NEXT")

    def title_dates_line(self, title, dates):
        self.set_font("Helvetica", "I", 9)
        self.set_text_color(*BLACK)
        self.cell(self.w - self.l_margin - self.r_margin - 40, 5, sanitize(title))
        self.set_font("Helvetica", "", 9)
        self.set_text_color(*GRAY)
        self.cell(40, 5, sanitize(dates), align="R", new_x="LMARGIN", new_y="NEXT")
        self.ln(1)

    def skill_line(self, category, skills):
        self.set_font("Helvetica", "", 9)
        self.set_text_color(*BLACK)
        self.set_x(self.l_margin + 5)
        self.cell(3, 4.5, "-")
        self.set_font("Helvetica", "B", 9)
        self.write(4.5, sanitize(f"{category}: "))
        self.set_font("Helvetica", "", 9)
        self.write(4.5, sanitize(skills))
        self.ln(5)


def _generate_experience_pdf(pdf, exp):
    if exp.get("company"):
        if exp.get("location"):
            pdf.company_line(exp["company"], exp["location"])
        else:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*DARK_BLUE)
            pdf.cell(0, 6, sanitize(exp["company"]), new_x="LMARGIN", new_y="NEXT")
    if exp.get("title"):
        pdf.title_dates_line(exp["title"], exp["dates"])
    else:
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*GRAY)
        pdf.cell(0, 5, sanitize(exp["dates"]), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

    if exp.get("intro"):
        pdf.body_text(exp["intro"])

    for sub_title, bullets in exp.get("subsections", []):
        if sub_title:
            pdf.sub_heading(sub_title)
        for bullet in bullets:
            pdf.bullet_point(bullet)
    pdf.ln(1)


def _render_simple_list(pdf, items):
    """Render a simple bulleted list in the PDF."""
    for item in items:
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*BLACK)
        pdf.set_x(pdf.l_margin + 5)
        pdf.cell(3, 4.5, "-")
        pdf.write(4.5, sanitize(item))
        pdf.ln(5)
    pdf.ln(3)


def generate_pdf(data, output_path):
    """Generate a PDF resume from structured data.

    Args:
        data: dict with keys: name, phone, email, linkedin, location,
              summary, skills (dict), experiences (list of dicts),
              education (list), certificates (list), awards (list), patents (list)
        output_path: path for the output .pdf file

    Returns:
        output_path
    """
    pdf = ResumePDF(candidate_name=data["name"])
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(20, 15, 20)

    # Name
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(*DARK_BLUE)
    pdf.cell(0, 12, sanitize(data["name"]), align="C", new_x="LMARGIN", new_y="NEXT")

    # Contact
    contact_parts = [data.get("phone", ""), data.get("email", ""),
                     data.get("linkedin", ""), data.get("location", "")]
    contact_line = "  ".join(p for p in contact_parts if p)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*GRAY)
    pdf.cell(0, 6, sanitize(contact_line), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # Summary
    pdf.section_heading("Summary")
    pdf.body_text(data["summary"])

    # Skills
    pdf.section_heading("Skills")
    for category, skills in data["skills"].items():
        pdf.skill_line(category, skills)
    pdf.ln(3)

    # Experience
    pdf.section_heading("Experience")
    for exp in data["experiences"]:
        _generate_experience_pdf(pdf, exp)

    # Education
    pdf.section_heading("Education")
    _render_simple_list(pdf, data.get("education", []))

    # Certificates
    if data.get("certificates"):
        pdf.section_heading("Certificates")
        _render_simple_list(pdf, data["certificates"])

    # Awards
    if data.get("awards"):
        pdf.section_heading("Awards")
        _render_simple_list(pdf, data["awards"])

    # Patents
    if data.get("patents"):
        pdf.section_heading("Patents")
        _render_simple_list(pdf, data["patents"])

    pdf.output(output_path)
    print(f"PDF resume saved to: {output_path}")
    return output_path
